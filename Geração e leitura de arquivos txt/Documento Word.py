from docx import Document
from docx.shared import Inches

doc = Document()

disciplina = input("Qual a discplina? ")
p = doc.add_paragraph("Disciplina: ")
p.add_run(disciplina).bold = True

carga = input("Qual a carga horária? ")
doc.add_paragraph('Com carga horária de {} horas.'.format(carga), style='List Bullet')

tab = doc.add_table(rows=1, cols=2)
tab.style="Colorful List Accent 6"
cels = tab.rows[0].cells
cels[0].text = 'Número'
cels[1].text = "Assunto"

for numero in range(1, 5):
   dados = tab.add_row().cells
   assuntos = input("Qual o assunto {}? ".format(numero))
   dados[0].text = str(numero)
   dados[1].text = assuntos

doc.save('c:\\temp\\meudoc.docx')

